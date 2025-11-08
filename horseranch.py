import hashlib
import os
import json
import datetime
from docx import Document

# ===============================
# FILES
# ===============================
HORSES_FILE = "horses.json"
BARNS_FILE = "barns.json"
USERS_FILE = "users.json"

# ===============================
# USER
# ===============================
def hash_password(password):
    """Return SHA256 hash of the password."""
    return hashlib.sha256(password.encode()).hexdigest()

def load_users():
    """Load users from file."""
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, "r") as f:
            return json.load(f)
    return []

def save_users(users):
    """Save users to file."""
    with open(USERS_FILE, "w") as f:
        json.dump(users, f, indent=4)

def register_user():
    """Register a new user."""
    users = load_users()
    username = input("Enter new username: ").strip()
    if any(u["username"] == username for u in users):
        print("Username already exists.\n")
        return False
    password = input("Enter password: ").strip()
    password_confirm = input("Confirm password: ").strip()
    if password != password_confirm:
        print("Passwords do not match.\n")
        return False
    users.append({
        "username": username,
        "password": hash_password(password)
    })
    save_users(users)
    print(f"User '{username}' registered successfully.\n")
    return True

def login_user():
    global current_user
    users = load_users()
    username = input("Username: ").strip()
    password = input("Password: ").strip()
    hashed = hash_password(password)
    for u in users:
        if u["username"] == username and u["password"] == hashed:
            print(f"\nWelcome, {username}!\n")
            current_user = username  # <-- track logged-in user
            return True
    print("Invalid username or password.\n")
    return False

# ===============================
# DATA STORAGE
# ===============================
horses = []
barns = []

# ===============================
# UTILITIES
# ===============================
def clear_screen():
    """Clear the terminal screen (works on Windows, Linux, Mac)."""
    os.system('cls' if os.name == 'nt' else 'clear')

def print_horse_details(horse):
    """Print all details of a horse."""
    print(f"\nName: {horse['name']}")
    print(f"Breed: {horse['breed']}")
    print(f"Barn: {horse['barn']}, Stall: {horse['stall']}")
    print(f"Birthdate: {horse['birth_date']}")
    print(f"Breakfast hay: {horse['breakfast_hay']}")
    print(f"Lunch hay: {horse['lunch_hay']}")
    print(f"Dinner hay: {horse['dinner_hay']}")
    print(f"Allergies: {', '.join(horse['allergies']) if horse['allergies'] else 'None'}")
    print("-" * 40)

# ===============================
# SAVE FUNCTION
# ===============================
def save_data():
    # Save horses
    all_horses = []
    if os.path.exists(HORSES_FILE):
        with open(HORSES_FILE, "r") as f:
            all_horses = json.load(f)
        # Remove current user's horses
        all_horses = [h for h in all_horses if h.get("owner") != current_user]
    # Add current user's horses
    all_horses.extend(horses)
    with open(HORSES_FILE, "w") as f:
        json.dump(all_horses, f, default=str, indent=4)
    # Save barns
    all_barns = []
    if os.path.exists(BARNS_FILE):
        with open(BARNS_FILE, "r") as f:
            all_barns = json.load(f)
        all_barns = [b for b in all_barns if b.get("owner") != current_user]
    all_barns.extend(barns)
    with open(BARNS_FILE, "w") as f:
        json.dump(all_barns, f, indent=4)
    print("Data saved successfully.\n")

# ===============================
# HORSE MANAGEMENT
# ===============================
def new_horse():
    print("\n=== Add a New Horse ===")
    if not barns:
        print("No barns available. Please add a barn first.\n")
        return

    while True:
        try:
            horse_name = input("Enter horse name: ").strip()
            date = input("Enter birth date (YYYY-MM-DD): ").strip()
            date_horse = datetime.datetime.strptime(date, "%Y-%m-%d").date()
            breed = input("Enter horse breed: ").strip()
            breakfast = input("Enter breakfast hay type: ").strip()
            lunch = input("Enter lunch hay type: ").strip()
            dinner = input("Enter dinner hay type: ").strip()
            allergies = [a.strip().lower() for a in input("Enter allergies (comma separated): ").split(",") if a.strip()]
            break
        except ValueError:
            print("Invalid date format. Please try again.\n")

    horse = {
        "name": horse_name,
        "birth_date": date_horse,
        "breed": breed,
        "breakfast_hay": breakfast,
        "lunch_hay": lunch,
        "dinner_hay": dinner,
        "allergies": allergies,
        "barn": None,
        "stall": None,
        "owner": current_user
    }

    horses.append(horse)
    assign_horse_to_stall(horse)
    save_data()
    print(f"\nHorse '{horse_name}' added successfully.\n")

def view_horse():
    print("\n=== View Horses ===")
    print("1. Print all horses")
    print("2. View a specific horse")
    print("3. Go back\n")
    while True:
        try:
            choice = int(input("Enter your choice: "))
            if choice == 1:
                if not horses:
                    print("\nNo horses registered yet.\n")
                else:
                    for x in horses:
                        print_horse_details(x)

                    export = input("\nWould you like to export all horses to a .docx file? (y/n): ").strip().lower()
                    if export == 'y':
                        doc = Document()
                        today = datetime.date.today().strftime("%Y-%m-%d")
                        doc.add_heading(f"Horse List – Generated on {today}", level=1)
                        table = doc.add_table(rows=1, cols=8)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = "Name"
                        hdr_cells[1].text = "Breed"
                        hdr_cells[2].text = "Barn"
                        hdr_cells[3].text = "Stall"
                        hdr_cells[4].text = "Birthdate"
                        hdr_cells[5].text = "Breakfast"
                        hdr_cells[6].text = "Lunch"
                        hdr_cells[7].text = "Dinner"

                        for x in horses:
                            row_cells = table.add_row().cells
                            row_cells[0].text = x['name']
                            row_cells[1].text = x['breed']
                            row_cells[2].text = x['barn'] if x['barn'] else "Unassigned"
                            row_cells[3].text = str(x['stall']) if x['stall'] else "-"
                            row_cells[4].text = str(x['birth_date'])
                            row_cells[5].text = x['breakfast_hay']
                            row_cells[6].text = x['lunch_hay']
                            row_cells[7].text = x['dinner_hay']

                        doc.save("horse_list.docx")
                        print("\nHorse list successfully exported to 'horse_list.docx'.\n")
                    else:
                        print("\nExport skipped.\n")

            elif choice == 2:
                print("\nSearch by:")
                print("1. Name (partial matches allowed)")
                print("2. Barn and Stall")
                search_choice = input("Enter your choice: ").strip()
                if search_choice == "1":
                    name_query = input("Enter horse name: ").strip().lower()
                    found = False
                    for x in horses:
                        if name_query in x['name'].lower():
                            print_horse_details(x)
                            found = True
                    if not found:
                        print("\nNo horses found with that name.\n")
                elif search_choice == "2":
                    barn_name = input("Enter barn name: ").strip().lower()
                    stall_number = input("Enter stall number: ").strip()
                    try:
                        stall_number = int(stall_number)
                    except ValueError:
                        print("Invalid stall number.\n")
                        continue
                    matches = [h for h in horses if h["barn"] and h["barn"].lower() == barn_name and h["stall"] == stall_number]
                    if matches:
                        for h in matches:
                            print_horse_details(h)
                    else:
                        print("No horse found in that barn/stall.\n")
                else:
                    print("\nInvalid choice. Please select 1 or 2.\n")
            elif choice == 3:
                return
            else:
                print("\nInvalid choice. Please select 1–3.\n")
        except ValueError:
            print("\nInvalid input. Please enter a number.\n")

def remove_horse():
    if not horses:
        print("\nNo horses registered yet.\n")
        return

    name = input("\nEnter the name of the horse to remove: ").strip().lower()
    for horse in horses:
        if horse["name"].lower() == name:
            confirm = input(f"Are you sure you want to remove '{horse['name']}'? (y/n): ").strip().lower()
            if confirm == "y":
                # Remove from barn
                for barn in barns:
                    if horse in barn['horses']:
                        barn['horses'].remove(horse)
                horses.remove(horse)
                print(f"\nHorse '{horse['name']}' has been removed successfully.\n")
                save_data()
            else:
                print("\nRemoval cancelled.\n")
            return
    print(f"\nHorse '{name}' not found.\n")

def assign_horse_to_stall(horse):
    if not barns:
        print("\nNo barns available. Please add a barn first.\n")
        return

    # Show barns with empty stalls
    available_barns = [b for b in barns if b['stalls'] > len(b['horses'])]
    if not available_barns:
        print("\nAll barns are full. Cannot assign horse.\n")
        return

    print("\nAvailable barns with empty stalls:")
    for i, barn in enumerate(available_barns, start=1):
        print(f"{i}. {barn['barn_name']} (Empty stalls: {barn['stalls'] - len(barn['horses'])})")

    while True:
        choice = input("\nEnter barn name or number to assign the horse: ").strip()
        selected_barn = None

        if choice.isdigit():
            index = int(choice) - 1
            if 0 <= index < len(available_barns):
                selected_barn = available_barns[index]
        else:
            for barn in available_barns:
                if barn["barn_name"].lower() == choice.lower():
                    selected_barn = barn
                    break

        if selected_barn:
            # Assign next available stall
            occupied_stalls = [h.get("stall") for h in selected_barn["horses"]]
            for stall_number in range(1, selected_barn["stalls"] + 1):
                if stall_number not in occupied_stalls:
                    horse["barn"] = selected_barn["barn_name"]
                    horse["stall"] = stall_number
                    selected_barn["horses"].append(horse)
                    print(f"\nHorse '{horse['name']}' assigned to Barn '{horse['barn']}', Stall {horse['stall']}.\n")
                    return
        else:
            print("Invalid barn choice. Please try again.\n")

def edit_horse():
    if not horses:
        print("\nNo horses available to edit.\n")
        return

    print("\n=== Edit a Horse ===")
    for i, horse in enumerate(horses, start=1):
        barn_info = f"{horse['barn']} (Stall {horse['stall']})" if horse.get("barn") else "Unassigned"
        print(f"{i}. {horse['name']} - {barn_info}")

    while True:
        choice = input("Enter horse number to edit: ").strip()
        if choice.isdigit():
            index = int(choice) - 1
            if 0 <= index < len(horses):
                selected_horse = horses[index]
                break
        print("Invalid choice. Try again.\n")

    print(f"\nEditing '{selected_horse['name']}' (leave blank to keep current value):")

    new_name = input(f"New name [{selected_horse['name']}]: ").strip()
    if new_name:
        selected_horse['name'] = new_name

    new_date = input(f"New birth date [{selected_horse['birth_date']}] (YYYY-MM-DD): ").strip()
    if new_date:
        try:
            selected_horse['birth_date'] = datetime.datetime.strptime(new_date, "%Y-%m-%d").date()
        except ValueError:
            print("Invalid date format. Keeping previous date.")

    new_breed = input(f"New breed [{selected_horse['breed']}]: ").strip()
    if new_breed:
        selected_horse['breed'] = new_breed

    new_breakfast = input(f"New breakfast hay [{selected_horse['breakfast_hay']}]: ").strip()
    if new_breakfast:
        selected_horse['breakfast_hay'] = new_breakfast

    new_lunch = input(f"New lunch hay [{selected_horse['lunch_hay']}]: ").strip()
    if new_lunch:
        selected_horse['lunch_hay'] = new_lunch

    new_dinner = input(f"New dinner hay [{selected_horse['dinner_hay']}]: ").strip()
    if new_dinner:
        selected_horse['dinner_hay'] = new_dinner

    new_allergies = input(f"New allergies (comma-separated) [{', '.join(selected_horse['allergies']) if selected_horse['allergies'] else 'None'}]: ").strip()
    if new_allergies:
        selected_horse['allergies'] = [a.strip().lower() for a in new_allergies.split(",") if a.strip()]

    reassign = input("Do you want to reassign this horse to a different barn? (y/n): ").strip().lower()
    if reassign == 'y':
        for barn in barns:
            if selected_horse in barn['horses']:
                barn['horses'].remove(selected_horse)
        assign_horse_to_stall(selected_horse)

    save_data()
    print(f"\nHorse '{selected_horse['name']}' updated successfully.\n")

# ===============================
# BARN MANAGEMENT
# ===============================
def barn_management():
    print("\n=== Barn Management ===")
    print("1. Add a barn")
    print("2. View a barn")
    print("3. Edit a barn")
    print("4. Search for empty stalls")
    print("5. Remove a barn")
    print("6. Go back\n")

    while True:
        try:
            choice = int(input("Enter your choice: "))
            if choice == 1:
                add_barn()
            elif choice == 2:
                view_barn()
            elif choice == 3:
                edit_barn()
            elif choice == 4:
                search_empty_stalls()
            elif choice == 5:
                remove_barn()
            elif choice == 6:
                return
            else:
                print("Invalid choice. Please select 1–6.\n")
        except ValueError:
            print("Invalid input. Please enter a number.\n")

def add_barn():
    print("\nAdd a Barn\n" + "-" * 20)
    while True:
        barn_name = input("Enter barn name: ").strip()
        if any(barn["barn_name"].lower() == barn_name.lower() for barn in barns):
            print(f"A barn named '{barn_name}' already exists.\n")
            continue
        try:
            stalls = int(input("Enter number of stalls: "))
            if stalls <= 0:
                print("Number of stalls must be greater than zero.\n")
                continue
            break
        except ValueError:
            print("Invalid number format.\n")

    barn = {
        "barn_name": barn_name,
        "stalls": stalls,
        "horses": [],
        "owner": current_user
    }
    barns.append(barn)
    print(f"\nBarn '{barn_name}' added successfully with {stalls} stalls.\n")
    save_data()

def view_barn():
    print("\nView a Barn\n" + "-" * 20)
    if not barns:
        print("No barns available.\n")
        return

    print("Available barns:")
    for i, barn in enumerate(barns, start=1):
        print(f"{i}. {barn['barn_name']} (Stalls: {barn['stalls']})")

    choice = input("\nEnter barn name or number to view: ").strip()
    selected_barn = None
    if choice.isdigit():
        index = int(choice) - 1
        if 0 <= index < len(barns):
            selected_barn = barns[index]
    else:
        for barn in barns:
            if barn["barn_name"].lower() == choice.lower():
                selected_barn = barn
                break

    if not selected_barn:
        print("\nBarn not found.\n")
        return

    print(f"\nBarn: {selected_barn['barn_name']}")
    print(f"Total stalls: {selected_barn['stalls']}")
    print(f"Occupied stalls: {len(selected_barn['horses'])}")
    print(f"Available stalls: {selected_barn['stalls'] - len(selected_barn['horses'])}\n")

    if selected_barn["horses"]:
        print("Horses in this barn:")
        for horse in selected_barn["horses"]:
            print(f" - {horse['name']}")
    else:
        print("No horses currently assigned.\n")

def edit_barn():
    if not barns:
        print("\nNo barns available to edit.\n")
        return

    print("\n=== Edit Barn ===")
    for i, barn in enumerate(barns, start=1):
        print(f"{i}. {barn['barn_name']} (Stalls: {barn['stalls']})")

    choice = input("\nEnter barn name or number to edit: ").strip()
    selected_barn = None

    if choice.isdigit():
        index = int(choice) - 1
        if 0 <= index < len(barns):
            selected_barn = barns[index]
    else:
        for barn in barns:
            if barn["barn_name"].lower() == choice.lower():
                selected_barn = barn
                break

    if not selected_barn:
        print("\nBarn not found.\n")
        return

    print(f"\nEditing Barn: {selected_barn['barn_name']}")
    note = input("Enter a note or description for this barn (leave blank to skip): ").strip()
    if note:
        selected_barn["note"] = note
        print(f"Note updated for barn '{selected_barn['barn_name']}'.")
        save_data()
    else:
        print("No changes made.")

def search_empty_stalls():
    print("\nSearch for Empty Stalls\n" + "-" * 30)
    if not barns:
        print("No barns have been added yet.\n")
        return

    empty_found = False
    for barn in barns:
        total = barn["stalls"]
        occupied = len(barn["horses"])
        available = total - occupied
        if available > 0:
            empty_found = True
            print(f"\nBarn: {barn['barn_name']}")
            print(f"Total stalls: {total}")
            print(f"Occupied: {occupied}")
            print(f"Available: {available}")
            print("-" * 25)
    if not empty_found:
        print("\nNo empty stalls found in any barn.\n")

def remove_barn():
    print("\nRemove a Barn\n" + "-" * 30)
    if not barns:
        print("No barns available to remove.\n")
        return

    name = input("Enter barn name to remove: ").strip().lower()
    found = False
    for barn in barns:
        if barn["barn_name"].lower() == name:
            confirm = input(f"Are you sure you want to remove '{barn['barn_name']}'? (y/n): ").strip().lower()
            if confirm == "y":
                # Remove horses from this barn
                for horse in barn['horses']:
                    horse['barn'] = None
                    horse['stall'] = None
                barns.remove(barn)
                print(f"\nBarn '{barn['barn_name']}' removed successfully.\n")
                save_data()
            else:
                print("\nRemoval canceled.\n")
            found = True
            break
    if not found:
        print(f"\nBarn '{name}' not found.\n")

# ===============================
# LOAD DATA FUNCTION
# ===============================
def load_data():
    global horses, barns
    if os.path.exists(HORSES_FILE):
        with open(HORSES_FILE, "r") as f:
            all_horses = json.load(f)
            horses = [h for h in all_horses if h.get("owner") == current_user]
            for h in horses:
                if "birth_date" in h:
                    h["birth_date"] = datetime.datetime.strptime(h["birth_date"], "%Y-%m-%d").date()
    if os.path.exists(BARNS_FILE):
        with open(BARNS_FILE, "r") as f:
            all_barns = json.load(f)
            barns = [b for b in all_barns if b.get("owner") == current_user]

# ===============================
# MAIN MENU
# ===============================
def main():
    load_data()
    while True:
        print("\n=== Horse & Barn Management System ===")
        print("1. Add new horse")
        print("2. View horses")
        print("3. Edit a horse")
        print("4. Remove a horse")
        print("5. Manage barns")
        print("6. Assign a horse to a stall")
        print("7. Exit\n")
        try:
            choice = int(input("Enter your choice: "))
            if choice == 1:
                new_horse()
            elif choice == 2:
                view_horse()
            elif choice == 3:
                edit_horse()
            elif choice == 4:
                remove_horse()
            elif choice == 5:
                barn_management()
            elif choice == 6:
                if not horses:
                    print("\nNo horses available to assign.\n")
                else:
                    for i, horse in enumerate(horses, start=1):
                        barn_info = f"{horse['barn']} (Stall {horse['stall']})" if horse.get("barn") else "Unassigned"
                        print(f"{i}. {horse['name']} - {barn_info}")
                    while True:
                        horse_choice = input("Enter horse number to assign/reassign: ").strip()
                        if horse_choice.isdigit():
                            index = int(horse_choice) - 1
                            if 0 <= index < len(horses):
                                assign_horse_to_stall(horses[index])
                                save_data()
                                break
                        print("Invalid choice. Try again.\n")
            elif choice == 7:
                print("\nExiting program. Goodbye!\n")
                save_data()
                break
            else:
                print("Invalid choice. Please select 1–7.\n")
        except ValueError:
            print("Invalid input. Please enter a number.\n")

if __name__ == "__main__":
    print("=== Welcome to Horse & Barn Management System ===")
    while True:
        print("1. Login")
        print("2. Register")
        print("3. Exit")
        choice = input("Enter choice: ").strip()
        if choice == "1":
            if login_user():
                main()  # go to main menu
                break
        elif choice == "2":
            register_user()
        elif choice == "3":
            print("Goodbye!")
            break
        else:
            print("Invalid choice.\n")
